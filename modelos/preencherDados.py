import os
from docx import Document

def substituir_texto(paragrafos, dados_dict):
    for p in paragrafos:
        for run in p.runs:
            for tag, valor in dados_dict.items():
                if tag in run.text:
                    run.text = run.text.replace(tag, str(valor))

def preencher_tabelas(documento, dados_dict):
     for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_texto(celula.paragraphs, dados_dict)
                
def preencher_modelo(caminho_template, pasta_destino, nome_arquivo_saida, dados):

    doc = Document(caminho_template)    

    preencher_tabelas(doc, dados)
    
    substituir_texto(doc.paragraphs, dados)

    # Garante que a pasta específica existe
    if not os.path.exists(pasta_destino):
        os.makedirs(pasta_destino)
        
    caminho_final = os.path.join(pasta_destino, nome_arquivo_saida.replace("/", "∕"))
    doc.save(caminho_final)
    print(f"Documento salvo com sucesso em: {caminho_final}")
